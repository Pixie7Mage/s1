"""Generate Word financial planning report from client payload."""

from __future__ import annotations

import io
import os
from datetime import date

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from client.calculations.helpers import (
    emergency_fund_required,
    format_inr,
    recommended_health_cover,
    recommended_term_cover,
    total_annual_expenses,
    total_annual_income,
)

def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Mulish"


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Mulish"


def _table_row(table, *values: str) -> None:
    row = table.add_row()
    for idx, val in enumerate(values):
        if idx < len(row.cells):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.text = ""
            run = p.add_run(val)
            run.font.name = "Mulish"
            run.font.size = Pt(10)


def _style_table_headers(row) -> None:
    for cell in row.cells:
        p = cell.paragraphs[0]
        text = p.text
        p.text = ""
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Mulish"
        run.font.size = Pt(10)


def _set_table_borders(table) -> None:
    tblPr = table._element.xpath('w:tblPr')
    if not tblPr:
        return
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')
        tblBorders.append(border)
    tblPr[0].append(tblBorders)


def _clear_table_borders(table) -> None:
    tblPr = table._element.xpath('w:tblPr')
    if not tblPr:
        return
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr[0].append(tblBorders)


def _set_cell_background(cell, hex_color: str) -> None:
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _num(value) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def _val(d: dict, camel: str, snake: str) -> float:
    return _num(d.get(camel, d.get(snake, 0)))


def _sum_extra(list_data: list) -> float:
    if not isinstance(list_data, list):
        return 0.0
    total = 0.0
    for item in list_data:
        if isinstance(item, dict):
            total += _num(item.get("amount", item.get("value", 0)))
    return total


def _add_pie_chart(doc: Document, labels: list[str], values: list[float], title: str) -> bool:
    non_zero = [(l, v) for l, v in zip(labels, values) if _num(v) > 0]
    if not non_zero:
        return False
    
    lbls, vals = zip(*non_zero)
    total = sum(vals)
    
    # Attractive, modern colors matching a premium brand palette
    colors = [
        '#0d9488', # Teal
        '#4f46e5', # Indigo
        '#f43f5e', # Rose
        '#d97706', # Amber
        '#10b981', # Emerald
        '#8b5cf6', # Purple
        '#06b6d4', # Sky
        '#ea580c'  # Orange
    ]
    
    fig, ax = plt.subplots(figsize=(7.2, 4.0), dpi=150)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    
    # Display percentages inside wedges only if they are > 4% to avoid crowding
    def my_autopct(pct):
        return f'{pct:.1f}%' if pct > 4 else ''
    
    wedges, texts, autotexts = ax.pie(
        vals,
        autopct=my_autopct,
        pctdistance=0.7,
        startangle=140,
        colors=colors[:len(vals)],
        wedgeprops=dict(edgecolor='white', linewidth=1.5, antialiased=True)
    )
    
    plt.setp(autotexts, size=9, weight="bold", color="white")
    
    # Build cleanly formatted labels for the legend
    legend_labels = []
    for l, v in zip(lbls, vals):
        pct = (v / total) * 100
        val_str = format_inr(v)
        legend_labels.append(f"{l}: {val_str} ({pct:.1f}%)")
    
    # Position the legend on the right of the pie chart
    ax.legend(
        wedges,
        legend_labels,
        title="Categories",
        loc="center left",
        bbox_to_anchor=(0.95, 0.5),
        frameon=True,
        facecolor='#f8fafc',
        edgecolor='#e2e8f0',
        fontsize=9
    )
    
    ax.set_title(title, fontsize=11, fontweight="bold", pad=15)
    plt.tight_layout()
    
    img_stream = io.BytesIO()
    # Save with bbox_inches='tight' to prevent any part of the legend or title from being cut off
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Insert the picture at a width of 5.5 inches for clean presentation
    run.add_picture(img_stream, width=Inches(5.5))
    return True


def _add_bar_chart(doc: Document, labels: list[str], current_values: list[float], future_values: list[float], title: str) -> bool:
    if not any(v > 0 for v in current_values):
        return False
        
    x = np.arange(len(labels))
    width = 0.35
    
    fig, ax = plt.subplots(figsize=(7.2, 4.0), dpi=150)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    
    ax.bar(x - width/2, current_values, width, label='Current Value', color='#1e3a5f')
    ax.bar(x + width/2, future_values, width, label='Value at Retirement', color='#0d9488')
    
    ax.set_ylabel('Amount (₹)', fontsize=10, fontweight='bold')
    ax.set_title(title, fontsize=11, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.legend(frameon=True, facecolor='#f8fafc', edgecolor='#e2e8f0', fontsize=9)
    
    def format_y_ticks(tick_val, pos):
        if tick_val >= 10_000_000:
            return f'{tick_val / 10_000_000:.1f} Cr'
        elif tick_val >= 100_000:
            return f'{tick_val / 100_000:.1f} L'
        return f'{tick_val:,.0f}'
        
    ax.yaxis.set_major_formatter(matplotlib.ticker.FuncFormatter(format_y_ticks))
    
    plt.tight_layout()
    
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_stream, width=Inches(5.5))
    return True


def calculate_sip(gap: float, rate: float, years: float) -> float:
    if years <= 0 or rate <= 0:
        return gap / (years * 12) if years > 0 else gap
    r_monthly = rate / 12
    months = years * 12
    factor = ((1 + r_monthly) ** months - 1) / r_monthly
    return gap / factor


def calculate_lumpsum(gap: float, rate: float, years: float) -> float:
    if years <= 0:
        return gap
    return gap / ((1 + rate) ** years)


def get_expected_return(years: float, equity_rate: float, debt_rate: float) -> float:
    if years <= 3:
        return debt_rate
    elif years <= 7:
        return (equity_rate + debt_rate) / 2
    else:
        return equity_rate


def generate_report(data: dict) -> bytes:
    doc = Document()

    # COVER PAGE (Simple White Page)
    title_p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("Financial Plan Report")
    run_title.bold = True
    run_title.font.size = Pt(24)
    run_title.font.name = "Mulish"

    # Add space to push down content
    for _ in range(5):
        doc.add_paragraph()

    personal = data.get("personal", {})
    name = personal.get("full_name", "Client")

    p_prepared = doc.add_paragraph()
    p_prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_prepared = p_prepared.add_run(f"Prepared for\n{name}")
    run_prepared.bold = True
    run_prepared.font.size = Pt(14)
    run_prepared.font.name = "Mulish"

    for _ in range(2):
        doc.add_paragraph()

    p_savart = doc.add_paragraph()
    p_savart.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_savart = p_savart.add_run("By\nSAVART ONE")
    run_savart.bold = True
    run_savart.font.size = Pt(12)
    run_savart.font.name = "Mulish"

    for _ in range(2):
        doc.add_paragraph()

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"Date\n{date.today().strftime('%d %B %Y')}")
    run_date.font.size = Pt(11)
    run_date.font.name = "Mulish"

    doc.add_page_break()

    # INDEX PAGE
    index_title = doc.add_heading("Index", level=1)
    index_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in index_title.runs:
        r.font.name = "Mulish"
    doc.add_paragraph()

    index_table = doc.add_table(rows=1, cols=2)
    _set_table_borders(index_table)
    hdr_cells = index_table.rows[0].cells
    hdr_cells[0].text = "Sr"
    hdr_cells[1].text = "Topics"
    _style_table_headers(index_table.rows[0])

    sections = [
        ("1", "Scope of the Report, Key Elements & Assumptions"),
        ("2", "Personal & Family Details"),
        ("3", "Income Analysis"),
        ("4", "Expenses Analysis"),
        ("5", "50/30/20 Guideline For Tracking Budget"),
        ("6", "Net Worth (Financial Assets, Fixed Assets & Liabilities)"),
        ("7", "Asset Allocation"),
        ("8", "Emergency Planning"),
        ("9", "Goals Summary & Recommendation"),
        ("10", "Investment Growth Projection at Retirement"),
        ("11", "Insurance Planning"),
        ("12", "Disclaimer"),
    ]

    for s_no, name_content in sections:
        _table_row(index_table, s_no, name_content)

    doc.add_page_break()

    # SECTION 1: Scope of the Report, Key Elements & Assumptions
    _heading(doc, "1. Scope of the Report, Key Elements & Assumptions")
    
    scope_text_pdf = (
        "A financial plan is a comprehensive evaluation of your current worth and "
        "future financial state by using current known variables to predict future "
        "income, asset values and withdrawal plans. This plan allocates your income "
        "to various types of expenses, such as rent or utilities, and reserves some "
        "income for short-term and long-term savings. And these savings are used "
        "to fulfill your future financial goals such as vacation, children education, "
        "retirement etc."
    )
    _para(doc, scope_text_pdf)
    doc.add_paragraph()

    scope_text_custom = (
        "This financial plan provides a comprehensive assessment of your current financial position and "
        "future financial outlook based on the information provided. It analyzes your income, expenses, "
        "assets, liabilities, and savings to create a structured roadmap for achieving your financial goals. "
        "The report recommends strategies for managing cash flow, building savings, optimizing investments, "
        "and planning for future objectives such as purchasing a home, children's education, vacations, and "
        "retirement, while supporting long-term financial security.This financial plan serves as a strategic "
        "decision-making tool and should be reviewed periodically to reflect changes in your financial "
        "circumstances, market conditions, personal goals, and applicable regulations."
    )
    _para(doc, scope_text_custom)
    doc.add_paragraph()

    _heading(doc, "Key Elements", level=2)
    key_el_table = doc.add_table(rows=1, cols=3)
    _set_table_borders(key_el_table)
    
    hdr_key = key_el_table.rows[0].cells
    hdr_key[0].text = "Sr"
    hdr_key[1].text = "Items"
    hdr_key[2].text = "Description"
    _style_table_headers(key_el_table.rows[0])

    key_elements = [
        ("1", "Financial Goals", "A financial plan is based on an individual's or a family's clearly defined financial goals."),
        ("2", "Net Worth Statement", "A snapshot of assets and liabilities serves as a benchmark for measuring progress towards financial goals."),
        ("3", "Cash Flow Analysis", "An income and spending plan determines how much can be set aside for debt repayment, savings and investing each month."),
        ("4", "Retirement Planning", "The plan should include a strategy for achieving retirement independent of other financial priorities."),
        ("5", "Insurance Planning", "Identify all risk exposures and provide the necessary coverage to protect the family and its assets against financial loss."),
        ("6", "Asset Allocation", "Include a customized asset allocation strategy based on specific investment objectives and a risk profile."),
    ]
    for sr, item, desc in key_elements:
        _table_row(key_el_table, sr, item, desc)

    doc.add_paragraph()
    _heading(doc, "Assumptions", level=2)
    assumptions = data.get("assumptions", {})
    t_assump = doc.add_table(rows=1, cols=2)
    _set_table_borders(t_assump)
    hdr_assump = t_assump.rows[0].cells
    hdr_assump[0].text = "Description"
    hdr_assump[1].text = "Value"
    _style_table_headers(t_assump.rows[0])

    _table_row(t_assump, "Inflation Rate", f"{assumptions.get('inflation', 7.0)} %")
    _table_row(t_assump, "Return on Debt Instruments", f"{assumptions.get('debt', 7.0)} %")
    _table_row(t_assump, "Return on Hybrid Funds", f"{(float(assumptions.get('equity', 12.0)) + float(assumptions.get('debt', 7.0))) / 2} %")
    _table_row(t_assump, "Return on Equity Instruments", f"{assumptions.get('equity', 12.0)} %")
    _table_row(t_assump, "Return on Gold", f"{assumptions.get('gold', 8.0)} %")
    _table_row(t_assump, "Emergency Fund Months", f"{assumptions.get('emergencyFundMonths', 6)} Months")
    _table_row(t_assump, "Life Expectancy", f"{assumptions.get('lifeExpectancy', 85)} Years")

    doc.add_page_break()

    # SECTION 2: Personal & Family Details
    _heading(doc, "2. Personal & Family Details")
    _para(doc, f"Name: {personal.get('full_name', '')}")
    _para(doc, f"Age: {personal.get('age', '')}")
    _para(doc, f"Occupation: {personal.get('occupation', '')}")
    doc.add_paragraph()

    family = data.get("family_members", [])
    if family:
        _para(doc, "Family Details:", bold=True)
        t_fam = doc.add_table(rows=1, cols=4)
        _set_table_borders(t_fam)
        hdr_fam = t_fam.rows[0].cells
        hdr_fam[0].text = "Name"
        hdr_fam[1].text = "Relationship"
        hdr_fam[2].text = "Age"
        hdr_fam[3].text = "Dependent"
        _style_table_headers(t_fam.rows[0])

        for member in family:
            _table_row(t_fam, member.get("name", ""), member.get("relationship", ""), str(member.get("age", "")), "Yes" if member.get("financially_dependent") else "No")
    else:
        _para(doc, "No family members recorded.")

    # Client Review
    client_review_text = assumptions.get("clientReview", assumptions.get("client_review", assumptions.get("riskProfile", assumptions.get("risk_profile", ""))))
    if client_review_text:
        doc.add_paragraph()
        _heading(doc, "Client Review", level=2)
        for line in client_review_text.split("\n"):
            line_stripped = line.strip()
            if line_stripped.startswith("-"):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line_stripped[1:].strip())
                run.font.name = "Mulish"
                run.font.size = Pt(11)
            elif line_stripped:
                p = doc.add_paragraph()
                run = p.add_run(line_stripped)
                run.bold = True
                run.font.name = "Mulish"
                run.font.size = Pt(11)

    doc.add_page_break()

    # SECTION 3: Income Analysis
    _heading(doc, "3. Income Analysis")
    income = data.get("income", {})
    total_income = income.get("total_annual") or total_annual_income(income)
    _para(doc, f"Total Income (Yearly): {format_inr(total_income)}", bold=True)
    doc.add_paragraph()

    inc_labels = ["Salary", "Business", "Rental", "Interest", "Other"]
    inc_values = [
        _num(income.get("salary", {}).get("annual")),
        _num(income.get("business", {}).get("annual")),
        _num(income.get("rental", {}).get("annual")),
        _num(income.get("interest", {}).get("annual")),
        _num(income.get("other", {}).get("annual")),
    ]
    for item in income.get("extra", []):
        inc_labels.append(item.get("name", "Extra"))
        inc_values.append(_num(item.get("annual", item.get("amount", 0))))
        
    _add_pie_chart(doc, inc_labels, inc_values, "Income Source Allocation")
    doc.add_paragraph()

    t_inc = doc.add_table(rows=1, cols=3)
    _set_table_borders(t_inc)
    hdr_inc = t_inc.rows[0].cells
    hdr_inc[0].text = "Income Type"
    hdr_inc[1].text = "Monthly (₹)"
    hdr_inc[2].text = "Yearly (₹)"
    _style_table_headers(t_inc.rows[0])

    for key, label in [
        ("salary", "Income from Salary"),
        ("business", "Business Income"),
        ("rental", "Rental Income"),
        ("interest", "Interest Income"),
        ("other", "Other Income"),
    ]:
        item = income.get(key, {})
        ann_val = _num(item.get("annual", item.get("amount", 0)))
        if ann_val > 0:
            _table_row(t_inc, label, format_inr(ann_val / 12), format_inr(ann_val))
            
    for item in income.get("extra", []):
        ann_val = _num(item.get("annual", item.get("amount", 0)))
        if ann_val > 0:
            _table_row(t_inc, item.get("name", "Other Income"), format_inr(ann_val / 12), format_inr(ann_val))
            
    _table_row(t_inc, "Total", format_inr(total_income / 12), format_inr(total_income))
    for cell in t_inc.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_page_break()

    # SECTION 4: Expenses Analysis
    _heading(doc, "4. Expenses Analysis")
    expenses = data.get("expenses", {})
    annual_exp = expenses.get("total_annual") or total_annual_expenses(expenses)
    _para(doc, f"Total Expenses (Yearly): {format_inr(annual_exp)}", bold=True)
    doc.add_paragraph()

    exp_labels = ["House Hold", "Lifestyle", "EMIs", "Other"]
    exp_values = [
        _num(expenses.get("monthly_household", 0)) * 12,
        _num(expenses.get("annual_lifestyle", 0)),
        _num(expenses.get("existing_emis", 0)),
        _num(expenses.get("other_expenses", 0)),
    ]
    for item in expenses.get("extra", []):
        exp_labels.append(item.get("name", "Extra"))
        exp_values.append(_num(item.get("annual", item.get("amount", 0))))
        
    _add_pie_chart(doc, exp_labels, exp_values, "Expenses Allocation")
    doc.add_paragraph()

    t_exp = doc.add_table(rows=1, cols=3)
    _set_table_borders(t_exp)
    hdr_exp = t_exp.rows[0].cells
    hdr_exp[0].text = "Expenses Type"
    hdr_exp[1].text = "Monthly (₹)"
    hdr_exp[2].text = "Yearly (₹)"
    _style_table_headers(t_exp.rows[0])

    expense_items = [
        ("monthly_household", "House Hold", True),
        ("annual_lifestyle", "Lifestyle", False),
        ("existing_emis", "Home Loan", False),
        ("other_expenses", "Other", False)
    ]
    for key, label, is_monthly in expense_items:
        val = _num(expenses.get(key, 0))
        yearly = val * 12 if is_monthly else val
        monthly = val if is_monthly else val / 12
        if yearly > 0:
            _table_row(t_exp, label, format_inr(monthly), format_inr(yearly))

    for item in expenses.get("extra", []):
        ann_val = _num(item.get("annual", item.get("amount", 0)))
        if ann_val > 0:
            _table_row(t_exp, item.get("name", "Other Expense"), format_inr(ann_val / 12), format_inr(ann_val))

    _table_row(t_exp, "Total", format_inr(annual_exp / 12), format_inr(annual_exp))
    for cell in t_exp.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_page_break()

    # SECTION 5: 50/30/20 Guideline For Tracking Budget
    _heading(doc, "5. 50/30/20 Guideline For Tracking Budget")
    
    t_guide = doc.add_table(rows=1, cols=3)
    _set_table_borders(t_guide)
    hdr_guide = t_guide.rows[0].cells
    hdr_guide[0].text = "EXPENSES: Maximum 50% Take Home Income"
    hdr_guide[1].text = "SAVING: Minimum 30%"
    hdr_guide[2].text = "FLOATING: 20%"
    _style_table_headers(t_guide.rows[0])

    row_guide = t_guide.add_row()
    for idx, text in enumerate([
        "- All household expenses : Utilities, food, entertainment, education, fuel.\n- Rent/House EMI",
        "- Saving for long term goals such as marriage, kids higher education, retirement",
        "- Short Term goals : Car, vacations, electronics\n- Emergency Fund : Medical, donations"
    ]):
        cell = row_guide.cells[idx]
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(text)
        run.font.name = "Mulish"
        run.font.size = Pt(10)

    doc.add_paragraph()

    savings = total_income - annual_exp
    savings_pct = (savings / total_income * 100) if total_income > 0 else 0
    expenses_pct = (annual_exp / total_income * 100) if total_income > 0 else 0

    _para(doc, f"Income = {format_inr(total_income)}")
    _para(doc, f"Expenses = {format_inr(annual_exp)} = {expenses_pct:.0f}% of Income")
    _para(doc, f"Floating = {format_inr(total_income * 0.20)} = 20% of Income")
    _para(doc, f"Saving = {format_inr(savings)} = {savings_pct:.0f}% of Income")
    doc.add_paragraph()

    if savings_pct >= 30:
        _para(doc, "You have a good Saving Ratio. Please make sure to sustain this in future as well.", bold=True)
    else:
        _para(doc, "Your Saving Ratio is below the recommended 30%. Consider optimizing your expenses to increase your savings.", bold=True)

    doc.add_page_break()

    # SECTION 6: Net Worth (Financial Assets, Fixed Assets & Liabilities)
    _heading(doc, "6. Net Worth (Financial Assets, Fixed Assets & Liabilities)")
    
    # 6.1 Financial Assets
    _heading(doc, "Financial Assets", level=2)
    assets = data.get("assets", {})
    
    t_fin_assets = doc.add_table(rows=1, cols=3)
    _set_table_borders(t_fin_assets)
    hdr_fin = t_fin_assets.rows[0].cells
    hdr_fin[0].text = "Financial Asset Type"
    hdr_fin[1].text = "Current Amount (₹)"
    hdr_fin[2].text = "Percentage"
    _style_table_headers(t_fin_assets.rows[0])

    # Predefined Financial map
    fin_map = [
        ("cash", "Cash"),
        ("savingsAccount", "Savings Account"),
        ("fd", "Fixed deposit (FD)"),
        ("stocks", "Equity / Stocks"),
        ("mutualFunds", "Mutual Fund"),
        ("esops", "ESOPs"),
        ("nps", "NPS (tier I & II)"),
        ("postalSavings", "Postal Savings"),
        ("epf", "EPF"),
        ("eps", "EPS"),
        ("amountReceivable", "Amount Receivable"),
        ("commodity", "Commodity (Gold/silver in biscuits, etc.)"),
    ]
    
    # Calculate totals
    total_predefined_fin = sum(_val(assets, k, k) for k, _ in fin_map)
    extra_fin = assets.get("extraFinancial", assets.get("extra_financial", []))
    if not isinstance(extra_fin, list):
        extra_fin = []
    total_extra_fin = _sum_extra(extra_fin)
    tot_fin_assets_val = total_predefined_fin + total_extra_fin

    for key, label in fin_map:
        val = _val(assets, key, key)
        if val > 0:
            pct = (val / tot_fin_assets_val * 100) if tot_fin_assets_val > 0 else 0
            _table_row(t_fin_assets, label, format_inr(val), f"{pct:.2f} %")
            
    for item in extra_fin:
        val = _num(item.get("amount", item.get("value", 0)))
        if val > 0:
            pct = (val / tot_fin_assets_val * 100) if tot_fin_assets_val > 0 else 0
            _table_row(t_fin_assets, item.get("name", "Custom Asset"), format_inr(val), f"{pct:.2f} %")

    _table_row(t_fin_assets, "Total Financial Assets", format_inr(tot_fin_assets_val), "100 %")
    for cell in t_fin_assets.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # 6.2 Fixed Assets
    _heading(doc, "Fixed Assets", level=2)
    t_fixed_assets = doc.add_table(rows=1, cols=3)
    _set_table_borders(t_fixed_assets)
    hdr_fix = t_fixed_assets.rows[0].cells
    hdr_fix[0].text = "Fixed Asset Type"
    hdr_fix[1].text = "Current Amount (₹)"
    hdr_fix[2].text = "Percentage"
    _style_table_headers(t_fixed_assets.rows[0])

    fixed_map = [
        ("realEstate", "Real Estate"),
        ("car", "Car"),
        ("bike", "Bike"),
        ("jewellery", "Jewellery (Gold/Silver)"),
        ("otherFixed", "Others"),
    ]
    
    total_predefined_fix = sum(_val(assets, k, k) for k, _ in fixed_map)
    # backward compat for older gold/realEstate keys if mapped
    if total_predefined_fix == 0:
        total_predefined_fix = _val(assets, "realEstate", "real_estate") + _val(assets, "gold", "gold")
        fixed_map = [("realEstate", "Real Estate"), ("gold", "Jewellery (Gold/Silver)")]

    extra_fix = assets.get("extraFixed", assets.get("extra_fixed", []))
    if not isinstance(extra_fix, list):
        extra_fix = []
    total_extra_fix = _sum_extra(extra_fix)
    tot_fixed_assets_val = total_predefined_fix + total_extra_fix

    for key, label in fixed_map:
        val = _val(assets, key, key)
        if val > 0:
            pct = (val / tot_fixed_assets_val * 100) if tot_fixed_assets_val > 0 else 0
            _table_row(t_fixed_assets, label, format_inr(val), f"{pct:.2f} %")
            
    for item in extra_fix:
        val = _num(item.get("amount", item.get("value", 0)))
        if val > 0:
            pct = (val / tot_fixed_assets_val * 100) if tot_fixed_assets_val > 0 else 0
            _table_row(t_fixed_assets, item.get("name", "Custom Asset"), format_inr(val), f"{pct:.2f} %")

    _table_row(t_fixed_assets, "Total Fixed Assets", format_inr(tot_fixed_assets_val), "100 %")
    for cell in t_fixed_assets.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # 6.3 Liabilities
    _heading(doc, "Liabilities", level=2)
    liabilities = data.get("liabilities", {})
    t_liab = doc.add_table(rows=1, cols=2)
    _set_table_borders(t_liab)
    hdr_liab = t_liab.rows[0].cells
    hdr_liab[0].text = "Liability Type"
    hdr_liab[1].text = "Current Outstanding (₹)"
    _style_table_headers(t_liab.rows[0])

    liab_map = [
        ("homeLoan", "Home Loan"),
        ("carLoan", "Car Loan / Auto Loan"),
        ("personalLoan", "Personal Loan"),
        ("goldLoan", "Gold Loan"),
        ("creditCard", "Credit Card Debt"),
        ("bankOverdraft", "Bank Overdraft"),
    ]
    
    total_predefined_liab = sum(_val(liabilities, k, k) for k, _ in liab_map)
    extra_liab = liabilities.get("extra", [])
    if not isinstance(extra_liab, list):
        extra_liab = []
    total_extra_liab = _sum_extra(extra_liab)
    tot_liab_val = total_predefined_liab + total_extra_liab

    for key, label in liab_map:
        val = _val(liabilities, key, key)
        if val > 0:
            _table_row(t_liab, label, format_inr(val))
            
    for item in extra_liab:
        val = _num(item.get("amount", item.get("value", 0)))
        if val > 0:
            _table_row(t_liab, item.get("name", "Custom Liability"), format_inr(val))

    _table_row(t_liab, "Total Liabilities", format_inr(tot_liab_val))
    for cell in t_liab.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()
    
    tot_assets_val = tot_fin_assets_val + tot_fixed_assets_val
    net_worth = tot_assets_val - tot_liab_val
    _heading(doc, f"Total Networth (Assets - Liabilities) = {format_inr(net_worth)}", level=2)

    doc.add_page_break()

    # SECTION 7: Asset Allocation
    _heading(doc, "7. Asset Allocation")
    
    # Calculate categorizations for allocation
    cash_val = _val(assets, "cash", "cash") + _val(assets, "savingsAccount", "savings_account")
    
    debt_val = (_val(assets, "fd", "fd") + _val(assets, "epf", "epf") + 
                _val(assets, "nps", "nps") + _val(assets, "postalSavings", "postal_savings") + 
                _val(assets, "eps", "eps"))
                
    equity_val = (_val(assets, "mutualFunds", "mutual_funds") + 
                  _val(assets, "stocks", "stocks") + 
                  _val(assets, "esops", "esops"))
                  
    gold_val = _val(assets, "gold", "gold") + _val(assets, "commodity", "commodity") + _val(assets, "jewellery", "jewellery")
    real_estate_val = _val(assets, "realEstate", "real_estate")
    
    # Add custom assets into cash/debt/equity allocations
    for item in extra_fin:
        name_lower = item.get("name", "").lower()
        val = _num(item.get("amount", item.get("value", 0)))
        if "cash" in name_lower or "saving" in name_lower:
            cash_val += val
        elif "equity" in name_lower or "stock" in name_lower or "fund" in name_lower:
            equity_val += val
        elif "gold" in name_lower or "silver" in name_lower or "commodity" in name_lower:
            gold_val += val
        else:
            debt_val += val # default custom financial to debt

    for item in extra_fix:
        name_lower = item.get("name", "").lower()
        val = _num(item.get("amount", item.get("value", 0)))
        if "real" in name_lower or "estate" in name_lower or "property" in name_lower or "land" in name_lower:
            real_estate_val += val
        elif "gold" in name_lower or "silver" in name_lower or "jewel" in name_lower:
            gold_val += val
        else:
            # other fixed assets (car, bike, etc.) can go to fixed/debt or we can just make it part of real estate/other
            real_estate_val += val
            
    alloc_labels = ["Cash", "Debt", "Equity", "Gold", "Real Estate"]
    alloc_values = [cash_val, debt_val, equity_val, gold_val, real_estate_val]
    
    _add_pie_chart(doc, alloc_labels, alloc_values, "Your Current Financial Asset Allocation")
    doc.add_paragraph()

    # Proposed Asset Allocation
    req_months = _num(assumptions.get("emergencyFundMonths", 6))
    req_ef = emergency_fund_required(annual_exp, req_months)
    
    total_fin = cash_val + debt_val + equity_val + gold_val + real_estate_val
    age = _num(personal.get("age", 35))
    
    prop_cash = min(total_fin, req_ef)
    rem_fin = max(0.0, total_fin - prop_cash)
    prop_equity = ((100 - age) / 100) * rem_fin
    prop_debt = rem_fin - prop_equity
    
    prop_labels = ["Cash (Emergency Fund)", "Debt", "Equity"]
    prop_values = [prop_cash, prop_debt, prop_equity]
    
    _add_pie_chart(doc, prop_labels, prop_values, "Proposed Financial Asset Allocation")

    doc.add_page_break()

    # SECTION 8: Emergency Planning
    _heading(doc, "8. Emergency Planning")
    
    emergency_text = (
        "Many a time, things don't pan out according to your plans. It makes sense then, to have a contingency plan "
        "armouring you to deal with unexpected outcomes. A contingency plan is often used to manage risk that "
        "may seem unlikely to happen, but if it does, would have a disruptive impact on your life.\n"
        "For example, a sudden job loss can derail your financial planning but most importantly, also leave you "
        "struggling to meet your everyday expenses and financial obligations like a personal or home loan till you "
        "find another job. Having a contingency fund will ensure that you don't have to worry if such a situation "
        "arises. It will enable you to pay your EMIs on time, even if your cash flow stops for some time as you have "
        "anticipated this and kept funds aside.\n"
        "Similarly, we need to have a rain cheque for medical emergencies or other natural disasters."
    )
    _para(doc, emergency_text)
    doc.add_paragraph()
    
    ef_data = data.get("emergency_fund", {})
    required_ef = _num(ef_data.get("required_fund", 0))
    ef_items = ef_data.get("items", [])
    if not isinstance(ef_items, list):
        ef_items = []
        
    assets = data.get("assets", {})
    available_ef = (
        _num(assets.get("savingsAccount", 0)) +
        _num(assets.get("cash", 0)) +
        _num(assets.get("sweepInFd", 0)) +
        _num(assets.get("liquidMutualFund", 0)) +
        _num(assets.get("moneyMarketFund", 0)) +
        _num(assets.get("overnightMutualFund", 0))
    )
    
    total_gap = max(0.0, required_ef - available_ef)
    
    _para(doc, f"Total Emergency Fund Available: {format_inr(available_ef)}")
    _para(doc, f"Total Emergency Fund Required: {format_inr(required_ef)}")
    _para(doc, f"Total Emergency Fund Gap: {format_inr(total_gap)}", bold=True)
    doc.add_paragraph()
    
    _heading(doc, "Manually Target Contingencies List", level=2)
    if not ef_items:
        _para(doc, "No emergency contingencies specified.")
    else:
        t_ef = doc.add_table(rows=1, cols=5)
        _set_table_borders(t_ef)
        headers_ef = ["Emergency Name", "Available (₹)", "Required (₹)", "Gap (₹)", "Where to Invest"]
        for i, h in enumerate(headers_ef):
            t_ef.rows[0].cells[i].text = h
        _style_table_headers(t_ef.rows[0])
        
        for item in ef_items:
            name = item.get("name", "")
            avail = _num(item.get("amount", 0))
            req = _num(item.get("required", 0))
            gap = max(0.0, req - avail)
            invest_to = item.get("where_to_invest", "")
            _table_row(t_ef, name, format_inr(avail), format_inr(req), format_inr(gap), invest_to)

    doc.add_page_break()

    # SECTION 9: Goals Summary & Recommendation
    _heading(doc, "9. Goals Summary & Recommendation")
    
    goals = data.get("goals", [])
    if not goals:
        _para(doc, "No goals specified.")
    else:
        t_goals = doc.add_table(rows=1, cols=4)
        _set_table_borders(t_goals)
        headers_goals = ["Goal Name", "Current Cost (₹)", "Target Corpus (₹)", "Interest (%)"]
        for i, h in enumerate(headers_goals):
            t_goals.rows[0].cells[i].text = h
        _style_table_headers(t_goals.rows[0])
                    
        for goal in goals:
            pres_cost = _num(goal.get("current_cost", 0))
            target_corp = _num(goal.get("target_corpus", 0))
            interest_rate = _num(goal.get("interest", 0))
            
            _table_row(t_goals, goal.get("name", ""), format_inr(pres_cost), format_inr(target_corp), f"{interest_rate}%")

    doc.add_page_break()

    # SECTION 10: Investment Growth Projection at Retirement
    personal = data.get("personal", {})
    age = _num(personal.get("age", 35))
    ret_age = _num(personal.get("retirement_age", 0))
    if ret_age <= 0:
        ret_age = max(60, age + 5)
    years = max(0, int(ret_age - age))
    
    assumptions = data.get("assumptions", {})
    eq_rate = _num(assumptions.get("equity", 12.0))
    d_rate = _num(assumptions.get("debt", 7.0))
    g_rate = _num(assumptions.get("gold", 8.0))
    inf_rate = _num(assumptions.get("inflation", 6.0))

    assets = data.get("assets", {})
    cash_val = _val(assets, "cash", "cash") + _val(assets, "savingsAccount", "savings_account")
    debt_val = (_val(assets, "fd", "fd") + _val(assets, "epf", "epf") + 
                _val(assets, "nps", "nps") + _val(assets, "postalSavings", "postal_savings") + 
                _val(assets, "eps", "eps"))
    equity_val = (_val(assets, "mutualFunds", "mutual_funds") + 
                  _val(assets, "stocks", "stocks") + 
                  _val(assets, "esops", "esops"))
    gold_val = _val(assets, "gold", "gold") + _val(assets, "commodity", "commodity") + _val(assets, "jewellery", "jewellery")
    real_estate_val = _val(assets, "realEstate", "real_estate")
    
    extra_fin = assets.get("extraFinancial", assets.get("extra_financial", []))
    if not isinstance(extra_fin, list):
        extra_fin = []
    for item in extra_fin:
        name_lower = item.get("name", "").lower()
        val = _num(item.get("amount", item.get("value", 0)))
        if "cash" in name_lower or "saving" in name_lower:
            cash_val += val
        elif "equity" in name_lower or "stock" in name_lower or "fund" in name_lower:
            equity_val += val
        elif "gold" in name_lower or "silver" in name_lower or "commodity" in name_lower:
            gold_val += val
        else:
            debt_val += val

    extra_fix = assets.get("extraFixed", assets.get("extra_fixed", []))
    if not isinstance(extra_fix, list):
        extra_fix = []
    for item in extra_fix:
        name_lower = item.get("name", "").lower()
        val = _num(item.get("amount", item.get("value", 0)))
        if "real" in name_lower or "estate" in name_lower or "property" in name_lower or "land" in name_lower:
            real_estate_val += val
        elif "gold" in name_lower or "silver" in name_lower or "jewel" in name_lower:
            gold_val += val
        else:
            real_estate_val += val

    fut_equity = equity_val * ((1 + eq_rate / 100) ** years)
    fut_debt = debt_val * ((1 + d_rate / 100) ** years)
    fut_gold = gold_val * ((1 + g_rate / 100) ** years)
    fut_cash = cash_val * ((1 + d_rate / 100) ** years)
    fut_real_estate = real_estate_val * ((1 + inf_rate / 100) ** years)
    
    total_current = equity_val + debt_val + gold_val + cash_val + real_estate_val
    total_future = fut_equity + fut_debt + fut_gold + fut_cash + fut_real_estate

    _heading(doc, "10. Investment Growth Projection at Retirement")
    
    intro_txt = (
        f"This section projects how your current financial investments will grow from your current age "
        f"of {age:.0f} to your targeted retirement age of {ret_age:.0f} (over {years} years). The calculations "
        f"compound your assets annually based on return rate assumptions specified in the Assumptions section."
    )
    _para(doc, intro_txt)
    doc.add_paragraph()
    
    if years > 0:
        labels_chart = ["Equity", "Debt", "Gold", "Cash/Savings", "Real Estate"]
        current_vals = [equity_val, debt_val, gold_val, cash_val, real_estate_val]
        future_vals = [fut_equity, fut_debt, fut_gold, fut_cash, fut_real_estate]
        _add_bar_chart(doc, labels_chart, current_vals, future_vals, "Investment Wealth Growth Projection")
        doc.add_paragraph()
        
        t_proj = doc.add_table(rows=1, cols=4)
        _set_table_borders(t_proj)
        hdr_proj = t_proj.rows[0].cells
        hdr_proj[0].text = "Asset Class"
        hdr_proj[1].text = "Current Value (₹)"
        hdr_proj[2].text = "Assumed Return (%)"
        hdr_proj[3].text = "Projected Value (₹)"
        _style_table_headers(t_proj.rows[0])
        
        _table_row(t_proj, "Equity (Mutual Funds, Stocks, ESOPs)", format_inr(equity_val), f"{eq_rate}%", format_inr(fut_equity))
        _table_row(t_proj, "Debt (FD, EPF, NPS, Postal)", format_inr(debt_val), f"{d_rate}%", format_inr(fut_debt))
        _table_row(t_proj, "Gold & Jewellery", format_inr(gold_val), f"{g_rate}%", format_inr(fut_gold))
        _table_row(t_proj, "Cash & Savings Account", format_inr(cash_val), f"{d_rate}%", format_inr(fut_cash))
        _table_row(t_proj, "Real Estate", format_inr(real_estate_val), f"{inf_rate}%", format_inr(fut_real_estate))
        
        _table_row(t_proj, "Total Assets", format_inr(total_current), "", format_inr(total_future))
        for cell in t_proj.rows[-1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
    else:
        _para(doc, "Growth calculations are skipped since years to retirement is 0.")

    doc.add_page_break()

    # SECTION 11: Insurance Planning
    _heading(doc, "11. Insurance Planning")
    
    insurance = data.get("insurance", {})
    policies = insurance.get("policies", [])
    if not isinstance(policies, list):
        policies = []
        
    if not policies:
        _para(doc, "No insurance policies added.")
    else:
        t_ins = doc.add_table(rows=1, cols=6)
        _set_table_borders(t_ins)
        hdr_ins = t_ins.rows[0].cells
        hdr_ins[0].text = "Insurance Type"
        hdr_ins[1].text = "Existing Cover"
        hdr_ins[2].text = "Recommended Cover"
        hdr_ins[3].text = "Premium Amount"
        hdr_ins[4].text = "Premium Tenure"
        hdr_ins[5].text = "Comment / Recommendation"
        _style_table_headers(t_ins.rows[0])
        
        for pol in policies:
            _table_row(
                t_ins,
                str(pol.get("policyType", "")),
                format_inr(float(pol.get("existingCover", 0.0) or 0.0)),
                format_inr(float(pol.get("recommendedCover", 0.0) or 0.0)),
                format_inr(float(pol.get("premium", 0.0) or 0.0)),
                str(pol.get("premiumTenure", "")),
                str(pol.get("comment", ""))
            )
        doc.add_paragraph()

    doc.add_page_break()

    # SECTION 12: Disclaimer
    _heading(doc, "12. Disclaimer")
    disclaimer_text = (
        "This financial plan was developed using information provided by you. Our Estimates of future returns and "
        "inflation parameters, using past history and reliable sources, play a significant part in the plan. While the "
        "information is presented in a detailed manner with exact numbers, PLEASE BE AWARE THAT ALL FUTURE "
        "PROJECTIONS ARE ESTIMATES ONLY.\n"
        "As the time period between the current date and projection date increases, so does the possible margin of "
        "error. This plan should be viewed as a 'road map', and it should be reviewed minimum every year and "
        "adjusted as more current or accurate information becomes available.\n"
        "All the calculations are done based on our proprietary algorithms and programs. As with other computer "
        "applications, these programs are subject to errors due to various reasons such as malware attack, hacking, "
        "human errors etc. Though we take the highest care to keep your information secured and making sure our "
        "algorithm works fine, still please do not consider this report as Final Financial Advice. There is human "
        "review, which has happened to this report but may prone to errors.\n"
        "Suggested Financial Plan to achieve your financial goals may not be accurate or yield expected results if "
        "the information provided by you is incorrect or any of the assumptions made are rendered invalid due to "
        "uncontrollable external forces, like change in interest rates, change in govt policies etc.\n"
        "This plan is not designed as a substitute for your own judgment, nor is it meant to eliminate the necessity of "
        "your personal review and analysis. This plan is designed to supplement your own planning and analysis to "
        "help you fulfil your financial objectives.Information provided in the attached report is general in nature and "
        "should NOT be construed as providing legal, accounting, investment and / or tax advice.Should you have "
        "any specific questions and / or issues in these areas, please consult your legal, tax, investment and / or "
        "accounting advisor. Savart One cannot endorse or support any specific decision you may make "
        "because we are not privy to all the other information that effective financial decision making requires."
    )
    _para(doc, disclaimer_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
