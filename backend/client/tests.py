import io
from docx import Document
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APITestCase

class GenerateReportViewTestCase(APITestCase):
    def test_generate_report_success(self):
        url = reverse('generate-report')
        payload = {
            "personal": {
                "full_name": "Jane Doe",
                "age": 35,
                "marital_status": "Married",
                "retirement_age": 60,
                "occupation": "Engineer"
            },
            "family_members": [],
            "income": {
                "salary": {"amount": 1000000, "period": "annual", "annual": 1000000}
            },
            "expenses": {
                "monthly_household": 30000,
                "annual_lifestyle": 100000,
                "existing_emis": 50000,
                "other_expenses": 20000
            },
            "assets": {
                "savingsAccount": 100000,
                "cash": 20000,
                "sweepInFd": 50000,
                "liquidMutualFund": 150000,
                "moneyMarketFund": 0,
                "overnightMutualFund": 0,
                "total": 320000
            },
            "liabilities": {
                "total": 1000000,
                "net_worth": 4000000
            },
            "goals": [
                {
                    "name": "Retirement",
                    "current_cost": 10000000,
                    "target_corpus": 20000000,
                    "interest": 6.0
                }
            ],
            "investments": {
                "mutual_funds": [],
                "stocks": []
            },
            "insurance": {
                "policies": [
                    {
                        "policyType": "Term Insurance",
                        "existingCover": 1000000,
                        "recommendedCover": 5000000,
                        "premium": 15000,
                        "premiumTenure": "20 Years",
                        "comment": "Good cover"
                    }
                ]
            },
            "emergency_fund": {
                "required_fund": 150000,
                "items": [
                    {
                        "name": "Medical Contingency",
                        "amount": 100000,
                        "required": 150000,
                        "gap": 50000,
                        "where_to_invest": "Savings Account"
                    }
                ]
            },
            "assumptions": {
                "inflation": 6,
                "equity": 12,
                "debt": 7,
                "gold": 8,
                "lifeExpectancy": 85,
                "emergencyFundMonths": 6
            }
        }
        
        response = self.client.post(url, payload, format='json')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(
            response['Content-Type'],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        self.assertIn("attachment; filename=", response['Content-Disposition'])
        
        # Load document from response content bytes
        doc_file = io.BytesIO(response.content)
        doc = Document(doc_file)
        
        # Check text in document paragraphs and tables
        all_text = []
        for p in doc.paragraphs:
            all_text.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        all_text.append(p.text)
            
        full_text = "\n".join(all_text)
        self.assertIn("Prepared for\nJane Doe", full_text)
        self.assertIn("By\nSAVART ONE", full_text)
        self.assertIn("Date\n", full_text)
        self.assertIn("1. Scope of the Report, Key Elements & Assumptions", full_text)
        self.assertIn("This financial plan provides a comprehensive assessment of your current financial position and future financial outlook based on the information provided.", full_text)
        
        # Let's check tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_texts.append(" | ".join(row_text))
                
        full_table_text = "\n".join(table_texts)
        self.assertIn("Sr | Topics", full_table_text)
        self.assertIn("1 | Scope of the Report, Key Elements & Assumptions", full_table_text)
        self.assertIn("2 | Personal & Family Details", full_table_text)
        self.assertIn("3 | Income Analysis", full_table_text)
        self.assertIn("10 | Investment Growth Projection at Retirement", full_table_text)
        self.assertIn("12 | Disclaimer", full_table_text)
